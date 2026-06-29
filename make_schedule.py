from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

BLUE_DARK  = RGBColor(0x1A, 0x52, 0x76)   # #1a5276
BLUE_MID   = RGBColor(0x29, 0x80, 0xB9)   # #2980b9
BLUE_LIGHT = RGBColor(0xEA, 0xF3, 0xFB)   # #eaf3fb
GRAY       = RGBColor(0x71, 0x80, 0x96)   # #718096
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x2D, 0x37, 0x48)

def set_cell_bg(cell, rgb_hex):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  rgb_hex)
    tcPr.append(shd)

def remove_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'nil')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_para_spacing(para, before=0, after=0):
    pPr  = para._p.get_or_add_pPr()
    spng = OxmlElement('w:spacing')
    spng.set(qn('w:before'), str(before))
    spng.set(qn('w:after'),  str(after))
    pPr.append(spng)

# ══════════════════════════════════════════════════════════════
# HEADER TABLE (blue banner)
# ══════════════════════════════════════════════════════════════
hdr_tbl = doc.add_table(rows=1, cols=1)
hdr_tbl.style = 'Table Grid'
hdr_cell = hdr_tbl.rows[0].cells[0]
set_cell_bg(hdr_cell, '1A5276')
remove_cell_borders(hdr_cell)
hdr_cell.width = Inches(6.5)

# Logo image centered in header
logo_path = '/root/.openclaw/workspace/vip-logo.jpg'
p_logo = hdr_cell.paragraphs[0]
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_logo, before=100, after=40)
if os.path.exists(logo_path):
    run_logo = p_logo.add_run()
    run_logo.add_picture(logo_path, width=Inches(2.0))

# Big welcome line
p_welcome = hdr_cell.add_paragraph()
p_welcome.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_welcome, before=40, after=0)
run = p_welcome.add_run('Welcome to the Team! 🎉')
run.font.size  = Pt(22)
run.font.bold  = True
run.font.color.rgb = WHITE

# Subtitle
p_sub = hdr_cell.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_sub, before=40, after=120)
run = p_sub.add_run('Training Schedule for ')
run.font.size  = Pt(12)
run.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
run2 = p_sub.add_run('[New Hire Name]')
run2.font.size  = Pt(12)
run2.font.bold  = True
run2.font.color.rgb = WHITE

doc.add_paragraph()  # spacer

# ══════════════════════════════════════════════════════════════
# GREETING
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
set_para_spacing(p, after=80)
r = p.add_run('Hi [Name],')
r.font.size = Pt(11)
r.font.bold = True
r.font.color.rgb = BLUE_DARK

p2 = doc.add_paragraph(
    "Welcome to VIP Medical Group! We're so excited to have you join our team and can't wait "
    "to get you started. Below is your training schedule for your first couple of weeks. You'll "
    "be working alongside our experienced sonographers to get comfortable with our workflow, "
    "protocols, and day-to-day rhythm."
)
p2.runs[0].font.size = Pt(11)
p2.runs[0].font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
set_para_spacing(p2, after=160)

# ══════════════════════════════════════════════════════════════
# SCHEDULE SECTION HEADER
# ══════════════════════════════════════════════════════════════
p_sh = doc.add_paragraph()
set_para_spacing(p_sh, before=80, after=80)
r = p_sh.add_run('📅  YOUR TRAINING SCHEDULE')
r.font.size  = Pt(9)
r.font.bold  = True
r.font.color.rgb = BLUE_MID

# ──────────────────────────────────────────────────────────────
# Helper: add a week block
# ──────────────────────────────────────────────────────────────
def add_week(doc, week_label, days):
    """
    days = list of (date_str, location, trainer)
    e.g. [('Tuesday, 2/17', 'Woodbridge', 'Julia'), ...]
    """
    # Week pill
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, 'EAF3FB')
    remove_cell_borders(cell)
    p = cell.paragraphs[0]
    set_para_spacing(p, before=60, after=60)
    r = p.add_run(f'  {week_label}  ')
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = BLUE_DARK
    doc.add_paragraph()  # tiny spacer

    # Day rows table
    day_tbl = doc.add_table(rows=len(days), cols=2)
    day_tbl.style = 'Table Grid'
    day_tbl.columns[0].width = Inches(2.6)
    day_tbl.columns[1].width = Inches(3.9)

    for i, (date_str, location, trainer) in enumerate(days):
        row = day_tbl.rows[i]
        # Date cell
        dc = row.cells[0]
        remove_cell_borders(dc)
        if i % 2 == 0:
            set_cell_bg(dc, 'F7FAFC')
        else:
            set_cell_bg(dc, 'FFFFFF')
        dp = dc.paragraphs[0]
        set_para_spacing(dp, before=80, after=80)
        dr = dp.add_run(f'  • {date_str}')
        dr.font.size = Pt(11)
        dr.font.bold = True
        dr.font.color.rgb = BLACK

        # Detail cell
        lc = row.cells[1]
        remove_cell_borders(lc)
        if i % 2 == 0:
            set_cell_bg(lc, 'F7FAFC')
        else:
            set_cell_bg(lc, 'FFFFFF')
        lp = lc.paragraphs[0]
        set_para_spacing(lp, before=80, after=80)
        lr1 = lp.add_run(location)
        lr1.font.size = Pt(11)
        lr1.font.bold = True
        lr1.font.color.rgb = BLUE_DARK
        lr2 = lp.add_run(f'  with {trainer}')
        lr2.font.size = Pt(11)
        lr2.font.color.rgb = GRAY

    doc.add_paragraph()  # spacer between weeks

# ══════════════════════════════════════════════════════════════
# WEEK 1
# ══════════════════════════════════════════════════════════════
add_week(doc, 'Week 1', [
    ('Tuesday, 2/17',   'Woodbridge',   'Julia'),
    ('Wednesday, 2/18', 'Princeton',    'Julia'),
    ('Thursday, 2/19',  'Woodbridge',   'Julia'),
    ('Friday, 2/20',    'Princeton',    'Julia'),
])

# WEEK 2
add_week(doc, 'Week 2', [
    ('Monday, 2/23',    'Woodland Park','Ashley'),
    ('Tuesday, 2/24',   'Woodbridge',   'Julia'),
    ('Wednesday, 2/25', 'Princeton',    'Julia'),
    ('Thursday, 2/26',  'Woodbridge',   'Julia'),
])

# ══════════════════════════════════════════════════════════════
# DIVIDER
# ══════════════════════════════════════════════════════════════
p_div = doc.add_paragraph('─' * 72)
p_div.runs[0].font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
set_para_spacing(p_div, before=80, after=80)

# ══════════════════════════════════════════════════════════════
# OFFICE ADDRESSES
# ══════════════════════════════════════════════════════════════
p_ah = doc.add_paragraph()
set_para_spacing(p_ah, before=80, after=100)
r = p_ah.add_run('📍  OFFICE ADDRESSES')
r.font.size = Pt(9)
r.font.bold = True
r.font.color.rgb = BLUE_MID

addr_tbl = doc.add_table(rows=1, cols=3)
addr_tbl.style = 'Table Grid'

offices = [
    ('Woodbridge',    '517 US-1 #1100\nIselin, NJ 08830'),
    ('Princeton',     '8 Forrestal Rd S, Suite 203\nPrinceton, NJ 08540'),
    ('Woodland Park', '1167 McBride Ave, Suite 2\nWoodland Park, NJ 07424'),
]

for i, (name, addr) in enumerate(offices):
    cell = addr_tbl.rows[0].cells[i]
    set_cell_bg(cell, 'F7FAFC')
    p = cell.paragraphs[0]
    set_para_spacing(p, before=80, after=20)
    r1 = p.add_run(name + '\n')
    r1.font.size = Pt(11)
    r1.font.bold = True
    r1.font.color.rgb = BLUE_DARK
    r2 = p.add_run(addr)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# CLOSING
# ══════════════════════════════════════════════════════════════
p_close = doc.add_paragraph(
    "Please plan to arrive a few minutes early each day so you have time to park and get settled. "
    "If anything comes up or you have questions before your first day, feel free to reach out — "
    "we're here to help!\n\n"
    "We're really looking forward to working with you and having you as part of the "
)
r_bold = p_close.add_run('VIP family')
r_bold.font.bold = True
r_bold.font.color.rgb = BLUE_DARK
p_close.add_run('. Welcome aboard! 💙')
for run in p_close.runs:
    run.font.size = Pt(11)
    if not run.font.color.type:
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
set_para_spacing(p_close, after=200)

# ══════════════════════════════════════════════════════════════
# FOOTER TABLE
# ══════════════════════════════════════════════════════════════
ftr_tbl = doc.add_table(rows=1, cols=1)
ftr_tbl.style = 'Table Grid'
ftr_cell = ftr_tbl.rows[0].cells[0]
set_cell_bg(ftr_cell, '1A5276')
remove_cell_borders(ftr_cell)

fp = ftr_cell.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(fp, before=100, after=20)
r = fp.add_run('VIP Medical Group')
r.font.size = Pt(11)
r.font.bold = True
r.font.color.rgb = WHITE

fp2 = ftr_cell.add_paragraph('Questions? Reply to this email or call us — we\'re happy to help.')
fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(fp2, before=0, after=100)
fp2.runs[0].font.size = Pt(10)
fp2.runs[0].font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)

doc.save('/root/.openclaw/workspace/VIP-Training-Schedule-Template.docx')
print('Done!')
